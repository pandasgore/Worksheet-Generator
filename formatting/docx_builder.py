from __future__ import annotations

import io
from pathlib import Path
import sys
from typing import Optional

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

if __package__ in (None, ""):
    PROJECT_ROOT = Path(__file__).resolve().parents[1]
    if str(PROJECT_ROOT) not in sys.path:
        sys.path.insert(0, str(PROJECT_ROOT))

from schemas.worksheet import WorksheetDocumentPayload, WorksheetDocumentProblem, DiagramSpecPayload
from formatting.text_cleaner import clean_math_text

# Import diagram generator (lazy to avoid import errors if matplotlib not installed)
_diagram_generator = None

def get_diagram_generator():
    """Lazy load the diagram generator."""
    global _diagram_generator
    if _diagram_generator is None:
        try:
            from formatting.diagram_generator import DiagramGenerator
            _diagram_generator = DiagramGenerator()
        except ImportError:
            print("Warning: Diagram generation not available (matplotlib not installed)")
            return None
    return _diagram_generator


def _friendly_difficulty_label(level: str) -> str:
    level = (level or "").lower()
    if level in {"easy", "medium", "hard"}:
        return level.capitalize()
    return "Mixed"


def _set_default_styles(doc: Document) -> None:
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(12)


def build_docx(
    worksheet: WorksheetDocumentPayload,
    output_path: str | Path,
    *,
    include_answer_key: bool = True,
) -> Path:
    document = Document()
    _set_default_styles(document)

    title = worksheet.title or f"Grade {worksheet.grade} Worksheet"
    concept = worksheet.concept
    difficulty = _friendly_difficulty_label(worksheet.difficulty)

    document.add_heading(title, level=1)

    name_line = document.add_paragraph()
    name_run = name_line.add_run("Name: ______________________    Date: ______________________")
    name_run.font.size = Pt(11)

    meta_parts = [
        f"Grade: {worksheet.grade}",
        f"Concept: {concept}",
        f"Difficulty: {difficulty}",
    ]
    document.add_paragraph(" • ".join(meta_parts))

    document.add_paragraph(
        "Show your work under each problem. Do your best and check your answers!"
    )
    document.add_paragraph("")

    if not worksheet.problems:
        document.add_paragraph("No problems were provided.")
    else:
        for idx, problem in enumerate(worksheet.problems, start=1):
            _add_problem(document, idx, problem)

    if include_answer_key and worksheet.problems:
        document.add_page_break()
        document.add_heading("Answer Key", level=1)
        document.add_paragraph(
            "For teacher reference. Print separately from the student copy."
        )
        for idx, problem in enumerate(worksheet.problems, start=1):
            _add_answer(document, idx, problem)

    output_path = Path(output_path)
    document.save(output_path)
    return output_path


def _add_diagram(document: Document, diagram_spec: DiagramSpecPayload) -> bool:
    """Add a diagram to the document. Returns True if successful."""
    generator = get_diagram_generator()
    if generator is None:
        return False
    
    try:
        from formatting.diagram_generator import DiagramSpec, DiagramType
        
        # Convert payload to DiagramSpec
        spec = DiagramSpec(
            diagram_type=DiagramType(diagram_spec.diagram_type),
            params=diagram_spec.params,
            width=diagram_spec.width,
            height=diagram_spec.height,
            show_labels=diagram_spec.show_labels,
            show_measurements=diagram_spec.show_measurements,
            title=diagram_spec.title,
        )
        
        # Generate the diagram as PNG bytes
        image_bytes = generator.generate(spec, format="png")
        
        # Add the image to the document
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(io.BytesIO(image_bytes), width=Inches(spec.width))
        
        return True
        
    except Exception as e:
        print(f"Warning: Could not generate diagram: {e}")
        return False


def _add_problem(document: Document, number: int, problem: WorksheetDocumentProblem) -> None:
    paragraph = document.add_paragraph(style="List Number")
    paragraph.add_run(f"Problem {number}: ").bold = True
    # Clean LaTeX notation from problem text
    cleaned_prompt = clean_math_text(problem.prompt.strip())
    paragraph.add_run(cleaned_prompt)

    # Add diagram if present (inline with the problem)
    if problem.diagram:
        document.add_paragraph("")  # Small gap before diagram
        _add_diagram(document, problem.diagram)

    # Add roughly 2 inches of whitespace (approx 8-10 empty lines)
    # Reduce if diagram was added
    num_lines = 5 if problem.diagram else 8
    for _ in range(num_lines):
        document.add_paragraph("")


def _add_answer(document: Document, number: int, problem: WorksheetDocumentProblem) -> None:
    header = document.add_paragraph()
    # Clean LaTeX notation from problem text in answer key
    cleaned_prompt = clean_math_text(problem.prompt.strip())
    header_run = header.add_run(f"{number}. {cleaned_prompt}")
    header_run.bold = True

    if problem.answer:
        ans = document.add_paragraph()
        # Clean LaTeX notation from answer text
        cleaned_answer = clean_math_text(problem.answer)
        ans_run = ans.add_run(f"Answer: {cleaned_answer}")
        ans_run.bold = True
    
    # Add answer diagram if present (for "create a diagram" type problems)
    if problem.answer_diagram:
        document.add_paragraph("")  # Small gap before diagram
        _add_diagram(document, problem.answer_diagram)

    if problem.solution_steps:
        document.add_paragraph("Steps:")
        for step in problem.solution_steps:
            # Clean LaTeX notation from solution steps
            cleaned_step = clean_math_text(step)
            document.add_paragraph(cleaned_step, style="List Bullet")


if __name__ == "__main__":
    sample_payload = WorksheetDocumentPayload(
        concept="Ratios",
        grade=6,
        difficulty="medium",
        num_problems=2,
        problems=[
            WorksheetDocumentProblem(prompt="A recipe uses 4 cups of flour and 6 cups of sugar. What is the ratio?", answer="2:3"),
            WorksheetDocumentProblem(prompt="Simplify the ratio 5:15.", answer="1:3"),
        ],
    )
    out = build_docx(sample_payload, "sample_output.docx")
    print(f"Wrote {out}")

